import os
import PyPDF2
from docx import Document
from typing import List, Optional, Tuple
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """文档处理器，支持PDF和Word文档的文本提取"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
    
    def extract_text(self, file_path: str) -> Tuple[str, List[str]]:
        """
        从文件中提取文本内容
        
        Args:
            file_path: 文件路径
            
        Returns:
            Tuple[str, List[str]]: (完整文本, 分块文本列表)
        """
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif file_ext == '.txt':
                return self._extract_from_txt(file_path)
            else:
                raise ValueError(f"不支持的文件格式: {file_ext}")
                
        except Exception as e:
            logger.error(f"文档处理失败: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> Tuple[str, List[str]]:
        """从PDF文件提取文本"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                full_text = ""
                text_chunks = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    full_text += page_text + "\n"
                    
                    # 分块处理（每页作为一个块）
                    if page_text.strip():
                        text_chunks.append(f"第{page_num + 1}页: {page_text.strip()}")
                
                return full_text.strip(), text_chunks
                
        except Exception as e:
            logger.error(f"PDF处理失败: {str(e)}")
            raise
    
    def _extract_from_docx(self, file_path: str) -> Tuple[str, List[str]]:
        """从Word文档提取文本"""
        try:
            doc = Document(file_path)
            full_text = ""
            text_chunks = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text += paragraph.text + "\n"
                    text_chunks.append(paragraph.text.strip())
            
            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        full_text += row_text + "\n"
                        text_chunks.append(row_text.strip())
            
            return full_text.strip(), text_chunks
            
        except Exception as e:
            logger.error(f"Word文档处理失败: {str(e)}")
            raise
    
    def _extract_from_txt(self, file_path: str) -> Tuple[str, List[str]]:
        """从文本文件提取文本"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                # 按段落分块
                chunks = [chunk.strip() for chunk in content.split('\n\n') if chunk.strip()]
                return content, chunks
                
        except Exception as e:
            logger.error(f"文本文件处理失败: {str(e)}")
            raise
    
    def chunk_text(self, text: str, chunk_size: int = 512, overlap: int = 50) -> List[str]:
        """
        将文本分块
        
        Args:
            text: 输入文本
            chunk_size: 块大小
            overlap: 重叠大小
            
        Returns:
            List[str]: 文本块列表
        """
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # 如果不是最后一块，尝试在句子边界分割
            if end < len(text):
                # 寻找最近的句号、问号或感叹号
                for i in range(end, max(start, end - 100), -1):
                    if text[i] in '.!?':
                        end = i + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
            if start >= len(text):
                break
        
        return chunks
    
    def clean_text(self, text: str) -> str:
        """
        清理文本内容
        
        Args:
            text: 原始文本
            
        Returns:
            str: 清理后的文本
        """
        # 移除多余的空白字符
        text = ' '.join(text.split())
        
        # 移除特殊字符（保留中文、英文、数字、标点）
        import re
        text = re.sub(r'[^\w\s\u4e00-\u9fff.,!?;:()\[\]{}"\'-]', '', text)
        
        return text.strip()
    
    def extract_resume_sections(self, text: str) -> dict:
        """
        从简历文本中提取各个章节
        
        Args:
            text: 简历文本
            
        Returns:
            dict: 各章节内容
        """
        sections = {
            'education': [],
            'experience': [],
            'skills': [],
            'projects': [],
            'achievements': []
        }
        
        # 简单的关键词匹配来识别章节
        lines = text.split('\n')
        current_section = None
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # 识别章节标题
            if any(keyword in line_lower for keyword in ['教育', 'education', '学历']):
                current_section = 'education'
            elif any(keyword in line_lower for keyword in ['经验', 'experience', '工作', '实习']):
                current_section = 'experience'
            elif any(keyword in line_lower for keyword in ['技能', 'skills', '技术']):
                current_section = 'skills'
            elif any(keyword in line_lower for keyword in ['项目', 'projects', '作品']):
                current_section = 'projects'
            elif any(keyword in line_lower for keyword in ['成就', 'achievements', '获奖']):
                current_section = 'achievements'
            elif line.strip() and current_section:
                sections[current_section].append(line.strip())
        
        return sections
